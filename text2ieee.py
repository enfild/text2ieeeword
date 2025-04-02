import openai
import docx
import os
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

openai.api_key = 'your_openai_api_key_here'

# open IEEE template
doc = docx.Document('ieee_template.docx')

# GPT-4 request
def format_with_gpt(text):
    response = openai.chat.completions.create(
        model='gpt-4o',
        messages=[
            {"role": "system", "content": "You are an assistant for formatting scientific articles according to the IEEE standard. Divide the text into sections: Abstract, Introduction, Methods, Results, Conclusion, References. It is clear which fragments belong to each section."},
            {"role": "user", "content": text}
        ],
        temperature=0.2
    )
    return response.choices[0].message.content.strip()

def insert_formatted_text(doc, formatted_text):
    sections = formatted_text.split("\n\n")
    for section in sections:
        lines = section.split("\n")
        header = lines[0].strip()
        content = "\n".join(lines[1:]).strip()

        p = doc.add_paragraph()
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        paragraphs = content.split("\n")
        for para in paragraphs:
            doc.add_paragraph(para.strip(), style='Normal')


with open('source_text.txt', 'r', encoding='utf-8') as file:
    original_text = file.read()


formatted_text = format_with_gpt(original_text)


insert_formatted_text(doc, formatted_text)

doc.save('formatted_ieee_paper.docx')

# unlock the document
os.chmod('formatted_ieee_paper.docx', 0o777)

print ("Formatted IEEE paper saved as 'formatted_ieee_paper.docx'")
